"""
Word document (.docx) text extractor.

Extracts text paragraph-by-paragraph using python-docx.
Tables are extracted row by row with cells joined by tabs.

Each logical section (paragraph or table) is returned as a page-like dict.
Since .docx files do not have physical pages, we group content into
approximate segments of ~1000 characters and label them as "section N".
"""

from pathlib import Path
from typing import Any

import docx

from utils.logger import get_logger

logger = get_logger(__name__)

_SECTION_SIZE = 1000  # characters per virtual "page" grouping


def extract(file_path: str | Path) -> list[dict[str, Any]]:
    """
    Extract text from a .docx file.

    Args:
        file_path: Path to the Word document.

    Returns:
        List of section-level dicts with text and metadata.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        logger.error("DOCX file not found: %s", file_path)
        return []

    raw_blocks: list[str] = []

    try:
        document = docx.Document(str(file_path))

        # Extract paragraphs
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                raw_blocks.append(text)

        # Extract tables
        for table in document.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    raw_blocks.append(row_text)

        logger.info(
            "Extracted %d text blocks from %s", len(raw_blocks), file_path.name
        )

    except Exception as exc:
        logger.error("Failed to extract DOCX %s: %s", file_path.name, exc)
        return []

    return _group_into_sections(raw_blocks, file_path.name)


def _group_into_sections(
    blocks: list[str], source_name: str
) -> list[dict[str, Any]]:
    """
    Group extracted text blocks into approximate page-sized sections.
    This creates a consistent metadata structure across document types.
    """
    sections: list[dict[str, Any]] = []
    current_text = ""
    section_index = 1

    for block in blocks:
        current_text += block + "\n"
        if len(current_text) >= _SECTION_SIZE:
            sections.append({
                "text":     current_text.strip(),
                "source":   source_name,
                "page":     section_index,
                "doc_type": "docx",
            })
            section_index += 1
            current_text = ""

    # Flush remaining content
    if current_text.strip():
        sections.append({
            "text":     current_text.strip(),
            "source":   source_name,
            "page":     section_index,
            "doc_type": "docx",
        })

    logger.info(
        "Grouped into %d sections for %s", len(sections), source_name
    )
    return sections